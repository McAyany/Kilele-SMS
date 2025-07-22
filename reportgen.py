import pandas as pd
from docx import Document
from docx2pdf import convert
import tkinter as tk
from tkinter import filedialog, messagebox
import os
import threading

cancel_flag = threading.Event()

def safe_val(val):
    return "-" if pd.isna(val) or val is None else str(val)

def get_marks(row_mid, row_end, col):
    m1 = safe_val(row_mid.get(col)) if row_mid is not None and col in row_mid.index else "-"
    m2 = safe_val(row_end.get(col)) if row_end is not None and col in row_end.index else "-"
    return m1, m2

def average_mark(val1, val2):
    nums = []
    for v in [val1, val2]:
        try:
            nums.append(float(v))
        except:
            continue
    return round(sum(nums) / len(nums), 1) if nums else "-"

def cbc_rating(average):
    try:
        avg = float(average)
    except:
        return "-"
    if avg >= 80: return "EE"
    elif avg >= 60: return "ME"
    elif avg >= 40: return "AE"
    return "BE"

def build_comment_map(df, subjects_info):
    comment_map = {}
    for subj, rating_col, comment_col in subjects_info:
        ratings_to_comments = {}
        if rating_col not in df.columns or comment_col not in df.columns:
            continue
        filtered = df[[rating_col, comment_col]].dropna()
        for _, row in filtered.iterrows():
            rating = str(row[rating_col]).strip()
            comment = str(row[comment_col]).strip()
            if rating and comment and rating not in ratings_to_comments:
                ratings_to_comments[rating] = comment
        comment_map[subj] = ratings_to_comments
    return comment_map

def replace_placeholders(doc, replacements):
    def process_runs(runs):
        full_text = ''.join(run.text for run in runs)
        for key, val in replacements.items():
            full_text = full_text.replace(key, str(val))
        for run in runs:
            run.text = ''
        if runs:
            runs[0].text = full_text

    for paragraph in doc.paragraphs:
        process_runs(paragraph.runs)
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    process_runs(paragraph.runs)

def generate_reports(excel_file, template_file, output_folder, on_done_callback):
    try:
        xls = pd.ExcelFile(excel_file)
        for sheet in ["MID TERM", "END TERM"]:
            if sheet not in xls.sheet_names:
                raise ValueError(f"Missing sheet: {sheet}")

        df_mid = xls.parse("MID TERM")
        df_end = xls.parse("END TERM")

        if 'STUDENT NAME' not in df_mid.columns or 'STUDENT NAME' not in df_end.columns:
            raise KeyError("Missing 'STUDENT NAME' column in one or both sheets.")

        all_names = set(df_mid['STUDENT NAME']).union(df_end['STUDENT NAME'])

        # Map: subject key in Excel -> template placeholders
        subject_map = {
            "MATHS": ("MATHS_MID", "MAT_END", "MR", "MC"),
            "ENG": ("ENG_MID", "ENG_END", "ER", "EC"),
            "KISW": ("KISW_MID", "KISW_END", "KR", "KC"),
            "INT-SCI": ("INT-SCI_MID", "INT-SCI_END", "IR", "IC"),
            "SSTRE": ("SSTRE_MID", "SSTRE_END", "SR", "SC"),
            "CREATIVE. ARTS": ("CREATIVE_MID", "CREATIVE_END", "CR", "CC"),
            "AGR/NUT": ("AGR_MID", "AGR_END", "AR", "AC")
        }

        # For extracting comments
        subjects_info = [(s, v[2], v[3]) for s, v in subject_map.items()]
        comment_map = build_comment_map(pd.concat([df_mid, df_end]), subjects_info)

        for name in all_names:
            if cancel_flag.is_set():
                print("Processing cancelled.")
                break

            row_mid = df_mid[df_mid['STUDENT NAME'] == name].iloc[0] if not df_mid[df_mid['STUDENT NAME'] == name].empty else None
            row_end = df_end[df_end['STUDENT NAME'] == name].iloc[0] if not df_end[df_end['STUDENT NAME'] == name].empty else None

            replacements = {"{{STUDENT NAME}}": name}

            for subj, (mid_ph, end_ph, rating_ph, comment_ph) in subject_map.items():
                m1, m2 = get_marks(row_mid, row_end, subj)
                avg = average_mark(m1, m2)
                rating = cbc_rating(avg)
                comment = comment_map.get(subj, {}).get(rating, "-")
                replacements[f"{{{{{mid_ph}}}}}"] = m1
                replacements[f"{{{{{end_ph}}}}}"] = m2
                replacements[f"{{{{{rating_ph}}}}}"] = rating
                replacements[f"{{{{{comment_ph}}}}}"] = comment

            for col in ['TOTAL', 'AVERAGE', 'PL', 'COMMENTS']:
                m1, m2 = get_marks(row_mid, row_end, col)
                val = m2 if m2 != '-' else m1
                replacements[f"{{{{{col}}}}}"] = val

            safe_name = "".join([c for c in name if c.isalnum() or c in (" ", "_", "-")]).strip().replace(" ", "_")
            docx_path = os.path.join(output_folder, f"{safe_name}_report.docx")
            pdf_path = os.path.join(output_folder, f"{safe_name}_report.pdf")

            doc = Document(template_file)
            replace_placeholders(doc, replacements)
            doc.save(docx_path)

            try:
                convert(docx_path, pdf_path)
            except Exception as e:
                print(f"Conversion error for {safe_name}: {e}")
            try:
                os.remove(docx_path)
            except:
                pass

        on_done_callback(cancelled=cancel_flag.is_set())
    except Exception as e:
        messagebox.showerror("Error", str(e))
        cancel_flag.clear()

def main_gui():
    root = tk.Tk()
    root.title("CBC PDF Report Generator")
    root.geometry("750x200")

    excel_var = tk.StringVar()
    template_var = tk.StringVar()
    output_var = tk.StringVar()

    tk.Label(root, text="Excel File:").grid(row=0, column=0, padx=10, pady=5)
    tk.Entry(root, textvariable=excel_var, width=50).grid(row=0, column=1)
    tk.Button(root, text="Browse", command=lambda: excel_var.set(filedialog.askopenfilename(filetypes=[("Excel files", "*.xlsx")]))).grid(row=0, column=2)

    tk.Label(root, text="Word Template:").grid(row=1, column=0, padx=10, pady=5)
    tk.Entry(root, textvariable=template_var, width=50).grid(row=1, column=1)
    tk.Button(root, text="Browse", command=lambda: template_var.set(filedialog.askopenfilename(filetypes=[("Word files", "*.docx")]))).grid(row=1, column=2)

    tk.Label(root, text="Output Folder:").grid(row=2, column=0, padx=10, pady=5)
    tk.Entry(root, textvariable=output_var, width=50).grid(row=2, column=1)
    tk.Button(root, text="Browse", command=lambda: output_var.set(filedialog.askdirectory())).grid(row=2, column=2)

    run_button = tk.Button(root, text="Generate Reports", bg="green", fg="white")
    cancel_button = tk.Button(root, text="Cancel", bg="red", fg="white")

    def on_done_callback(cancelled=False):
        run_button.config(state="normal")
        cancel_button.config(state="disabled")
        root.config(cursor="")
        if cancelled:
            messagebox.showinfo("Cancelled", "Report generation was cancelled.")
        else:
            messagebox.showinfo("Success", "PDF reports generated!")

    def run():
        if not (excel_var.get() and template_var.get() and output_var.get()):
            messagebox.showerror("Missing Input", "Please select all required files/folders.")
            return
        cancel_flag.clear()
        run_button.config(state="disabled")
        cancel_button.config(state="normal")
        root.config(cursor="watch")
        threading.Thread(
            target=generate_reports,
            args=(excel_var.get(), template_var.get(), output_var.get(), on_done_callback),
            daemon=True
        ).start()

    def cancel():
        cancel_flag.set()

    run_button.config(command=run)
    cancel_button.config(command=cancel)

    run_button.grid(row=3, column=1, pady=15)
    cancel_button.grid(row=3, column=2, pady=15)

    root.mainloop()

if __name__ == "__main__":
    main_gui()
